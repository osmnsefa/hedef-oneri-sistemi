import os
import sys

# Proje ana dizinini sys.path'e ekle (doğrudan çalıştırma desteği için)
base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if base_dir not in sys.path:
    sys.path.append(base_dir)

import pandas as pd
from docx import Document
import logging
from src.config import Config
from langchain_text_splitters import RecursiveCharacterTextSplitter
logger = logging.getLogger(__name__)

class DataLoader:
    def __init__(self):
        self.data_dir = Config.DATA_DIR
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ".", " ", ""]
        )
        
    def load_excel_data(self):
        """Excel dosyalarından her satırı ayrı bir döküman olarak çıkarır."""
        documents = []
        if not os.path.exists(self.data_dir):
            logger.warning(f"Veri klasörü bulunamadı: {self.data_dir}")
            return documents

        for filename in os.listdir(self.data_dir):
            if filename.endswith(('.xlsx', '.xls')) and not filename.startswith('~$'):
                file_path = os.path.join(self.data_dir, filename)
                try:
                    df = pd.read_excel(file_path).fillna("")
                    columns = df.columns.tolist()
                    
                    # Kolon isimlerini temizle (boşluk vs)
                    df.columns = [str(c).strip() for c in df.columns]
                    columns = df.columns.tolist()

                    for index, row in df.iterrows():
                        row_text = []
                        metadata_name = ""
                        
                        # İsim tespiti için daha esnek kontrol
                        name_col = next((c for c in columns if c.lower() in ['isim', 'ad soyad', 'çalışan']), None)
                        if name_col:
                            metadata_name = str(row[name_col]).strip()

                        for col in columns:
                            val = str(row[col]).strip()
                            if val and val.lower() != "nan":
                                row_text.append(f"{col}: {val}")
                        
                        if row_text:
                            content = " | ".join(row_text)
                            documents.append({
                                "source": filename, 
                                "content": content,
                                "employee": metadata_name,
                                "is_excel": True # Excel olduğunu işaretle
                            })
                    
                    logger.info(f"✅ Excel yüklendi: {filename} ({len(df)} satır)")
                except Exception as e:
                    logger.error(f"❌ Excel okuma hatası ({filename}): {str(e)}")
        return documents

    def load_word_data(self):
        """Word dosyalarından ham metin verisi çıkarır."""
        documents = []
        if not self.data_dir or not os.path.exists(self.data_dir):
            return documents

        for filename in os.listdir(self.data_dir):
            if filename.endswith('.docx'):
                file_path = os.path.join(self.data_dir, filename)
                try:
                    doc = Document(file_path)
                    content_parts = [f"--- DOSYA: {filename} ---"]
                    content_parts.append("\n".join([para.text for para in doc.paragraphs if para.text.strip()]))
                    
                    # Tabloları da oku
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_text:
                                content_parts.append(" | ".join(row_text))
                    
                    full_text = "\n".join(content_parts)
                    documents.append({"source": filename, "content": full_text})
                    logger.info(f"✅ Word yüklendi: {filename}")
                except Exception as e:
                    logger.error(f"❌ Word okuma hatası ({filename}): {str(e)}")
        return documents

    def get_chunked_documents(self):
        """Tüm verileri yükler ve RAG için hazırlar."""
        excel_docs = self.load_excel_data()
        word_docs = self.load_word_data()
        
        final_docs = []
        
        # Excel verilerini doğrudan (split etmeden) ekle
        for doc in excel_docs:
            final_docs.append({
                "page_content": doc['content'],
                "metadata": {
                    "source": doc['source'],
                    "employee": doc.get('employee', 'Bilinmeyen')
                }
            })
            
        # Word verilerini split ederek ekle
        for doc in word_docs:
            chunks = self.text_splitter.split_text(doc['content'])
            for chunk in chunks:
                final_docs.append({
                    "page_content": chunk,
                    "metadata": {"source": doc['source']}
                })
                
        logger.info(f"Toplam {len(final_docs)} döküman hazırlandı.")
        return final_docs

    def get_dropdown_options(self):
        """Excel'den benzersiz Çalışan isimlerini ve Hedef Türlerini çeker."""
        employees = set()
        target_types = set()
        
        if not os.path.exists(self.data_dir):
            return [], []

        for filename in os.listdir(self.data_dir):
            if filename.endswith(('.xlsx', '.xls')) and not filename.startswith('~$'):
                file_path = os.path.join(self.data_dir, filename)
                try:
                    df = pd.read_excel(file_path).fillna("")
                    
                    # Kolon adı düzeltmesi: 'İsim' veya 'Ad Soyad' olabilir
                    name_col = 'İsim' if 'İsim' in df.columns else 'Ad Soyad'
                    
                    if name_col in df.columns:
                        employees.update(df[name_col].dropna().astype(str).unique())
                    
                    if 'Hedef Türü' in df.columns:
                        target_types.update(df['Hedef Türü'].dropna().astype(str).unique())
                        
                except Exception as e:
                    logger.error(f"Metadata okuma hatası ({filename}): {str(e)}")
        
        return sorted(list(employees)), sorted(list(target_types))

    def get_employee_history(self, employee_name, target_type=None):
        """Seçilen çalışan ve hedef türü için geçmiş verileri tablo olarak döner."""
        history_df = pd.DataFrame()
        
        if not os.path.exists(self.data_dir):
            return history_df

        for filename in os.listdir(self.data_dir):
            if filename.endswith(('.xlsx', '.xls')) and not filename.startswith('~$'):
                file_path = os.path.join(self.data_dir, filename)
                try:
                    df = pd.read_excel(file_path).fillna("")
                    name_col = 'İsim' if 'İsim' in df.columns else 'Ad Soyad'
                    
                    if name_col in df.columns:
                        # Çalışana göre filtrele
                        filtered = df[df[name_col] == employee_name]
                        
                        # Hedef türüne göre filtrele (opsiyonel)
                        if target_type and 'Hedef Türü' in df.columns:
                            filtered = filtered[filtered['Hedef Türü'] == target_type]
                            
                        if not filtered.empty:
                            history_df = pd.concat([history_df, filtered], ignore_index=True)
                            
                except Exception as e:
                    logger.error(f"Geçmiş verisi okuma hatası ({filename}): {str(e)}")
                    
        return history_df

    def get_employee_metadata(self, employee_name):
        """Çalışanın kimlik bilgilerini (Unvan, Bölüm, Sicil) döner."""
        metadata = {}
        
        if not os.path.exists(self.data_dir):
            return metadata

        for filename in os.listdir(self.data_dir):
            if filename.endswith(('.xlsx', '.xls')) and not filename.startswith('~$'):
                file_path = os.path.join(self.data_dir, filename)
                try:
                    df = pd.read_excel(file_path).fillna("")
                    name_col = 'İsim' if 'İsim' in df.columns else 'Ad Soyad'
                    
                    if name_col in df.columns:
                        # Çalışana göre filtrele
                        person_row = df[df[name_col] == employee_name]
                        
                        if not person_row.empty:
                            row = person_row.iloc[0]
                            # İstenen sütunlar
                            target_cols = ['Sicil', 'Unvan', 'Bölüm Ana Sorumluluk Alanı']
                            for col in target_cols:
                                if col in df.columns:
                                    metadata[col] = row[col]
                            
                            # Bulduysak çıkalım (ilk eşleşme yeterli varsayımı)
                            if metadata:
                                return metadata
                            
                except Exception as e:
                    logger.error(f"Metadata okuma hatası ({filename}): {str(e)}")
                    
        return metadata