import os
import re
import logging
import pandas as pd
from datetime import datetime
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
import chromadb
from docx import Document
from chromadb.utils import embedding_functions
from langchain_text_splitters import RecursiveCharacterTextSplitter

from src.config import Config
from src.models import Base, User, Employee, PerformanceHistory, JobDescriptions, EmployeeFeedback
from werkzeug.security import generate_password_hash

logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────────────────────────────────────
# Sabit Eşlemeler
# ──────────────────────────────────────────────────────────────────────────────

# Geri Bildirimler.docx bölüm başlıkları -> models.py EmployeeFeedback kolon isimleri
FEEDBACK_COLUMN_MAP = {
    "güçlü alanlar": "guclu_alanlar",
    "gelişime açık alanlar": "gelisim_alanlari",
    "gelişim alanları": "gelisim_alanlari",
    "genel performans": "genel_degerlendirme",
    "yılından beklentiler": "gelecek_beklentileri",
    "beklenti": "gelecek_beklentileri"
}

# ORGANİZASYON ŞEMASI bölüm anahtar kelimeleri
ORG_POSITION_MARKER    = "görev tanımı"
ORG_RESPONSIBILITIES   = "ana sorumluluklar"
ORG_COMPETENCIES       = {"aranan yetkinlikler", "yetkinlikler"}
ORG_SKIP_SECTIONS      = {"pozisyonun amacı"}


# ──────────────────────────────────────────────────────────────────────────────
# Fonksiyon 1: Görev Tanımları Ayrıştırma
# ──────────────────────────────────────────────────────────────────────────────

def parse_job_descriptions(session, data_dir):
    """
    ORGANİZASYON ŞEMASI VE GÖREV TANIMLARI.docx dosyasını okur ve
    JobDescriptions tablosunu doldurur.

    Mantık:
    - "Görev Tanımı" içeren satır → yeni bir pozisyon başlığıdır.
    - "Ana Sorumluluklar" başlığından sonraki satırlar → responsibilities.
    - "Aranan Yetkinlikler" / "Yetkinlikler" başlığından sonraki satırlar → competencies.
    - Her çalıştırmada tabloyu önce temizler (idempotent).
    """
    word_path = os.path.join(data_dir, "ORGANİZASYON ŞEMASI VE GÖREV TANIMLARI.docx")
    if not os.path.exists(word_path):
        print(f"[UYARI] Dosya bulunamadı: {word_path}")
        return

    print("JobDescriptions ayrıştırılıyor...")

    # İdempotent: mevcut kayıtları temizle
    session.query(JobDescriptions).delete()
    session.commit()

    doc = Document(word_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    current_position = None
    current_section  = None   # 'responsibilities' | 'competencies' | None
    buff = {"responsibilities": [], "competencies": []}

    def flush_position():
        """Biriktirilen pozisyonu veritabanına yazar ve buffer'ı sıfırlar."""
        nonlocal current_position, buff
        if not current_position:
            return
        jd = JobDescriptions(
            position_name=current_position,
            responsibilities="\n".join(buff["responsibilities"]) or None,
            technical_requirements=None,   # Belge bu alanı ayrı tutmuyorsa None
            competencies="\n".join(buff["competencies"]) or None,
        )
        session.add(jd)
        current_position = None
        buff = {"responsibilities": [], "competencies": []}

    for line in paragraphs:
        lower = line.lower()

        # ── Yeni pozisyon başlığı ────────────────────────────────────
        if ORG_POSITION_MARKER in lower:
            flush_position()

            # "– Görev Tanımı", "- Görev Tanımı" gibi ekleri temizle
            name = line
            for suffix in ["– Görev Tanımı", "- Görev Tanımı", "Görev Tanımı"]:
                name = name.replace(suffix, "").strip()
            # "1.", "2.1." gibi numara öneklerini kaldır
            name = re.sub(r"^\d+(\.\d+)*\.?\s*", "", name).strip()
            current_position = name
            current_section  = None
            continue

        # ── Bölüm başlıkları ────────────────────────────────────────
        if lower == ORG_RESPONSIBILITIES:
            current_section = "responsibilities"
            continue

        if lower in ORG_COMPETENCIES:
            current_section = "competencies"
            continue

        if lower in ORG_SKIP_SECTIONS:
            current_section = None
            continue

        # ── İçerik satırı ───────────────────────────────────────────
        if current_position and current_section:
            buff[current_section].append(line)

    flush_position()   # Dosyanın sonundaki son pozisyonu kaydet
    session.commit()

    count = session.query(JobDescriptions).count()
    print(f"JobDescriptions: {count} pozisyon başarıyla kaydedildi.")


# ──────────────────────────────────────────────────────────────────────────────
# Fonksiyon 2: Geri Bildirimler Ayrıştırma
# ──────────────────────────────────────────────────────────────────────────────

import difflib

def parse_feedbacks(session, data_dir):
    """
    Geri Bildirimler.docx dosyasını State Machine (Durum Makinesi) mantığıyla okur.
    Hiçbir kural satır sırasına bağlı değildir; her satır kendi içeriğine göre değerlendirilir.
    """
    word_path = os.path.join(data_dir, "Geri Bildirimler.docx")
    if not os.path.exists(word_path):
        print(f"[UYARI] Dosya bulunamadı: {word_path}")
        return

    print("EmployeeFeedback ayrıştırılıyor...")
    session.query(EmployeeFeedback).delete()
    session.commit()

    doc = Document(word_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # 1. Sözlük (dict) Oluştur
    all_employees = session.query(Employee).all()
    names_dict = {}
    for emp in all_employees:
        ad_soyad = f"{emp.first_name} {emp.last_name}".lower()
        names_dict[ad_soyad] = emp.user_sicil

    known_names = list(names_dict.keys())

    # 2. State Variables (Durum Değişkenleri)
    current_sicil = None
    current_year = None
    current_category = None
    
    current_dict = {
        "harf_notu": "",
        "genel_degerlendirme": [],
        "guclu_alanlar": [],
        "gelisim_alanlari": [],
        "gelecek_beklentileri": []
    }
    
    def flush_buffer():
        nonlocal current_sicil, current_year, current_category, current_dict
        if current_sicil and current_year:
            # En az bir değer girilmiş mi kontrol et
            has_data = any((val for key, val in current_dict.items() if key != "harf_notu")) or current_dict["harf_notu"]
            if has_data:
                fb = EmployeeFeedback(
                    employee_sicil=current_sicil,
                    yil=int(current_year) if current_year.isdigit() else None,
                    harf_notu=current_dict["harf_notu"] or None,
                    genel_degerlendirme="\n".join(current_dict["genel_degerlendirme"]) or None,
                    guclu_alanlar="\n".join(current_dict["guclu_alanlar"]) or None,
                    gelisim_alanlari="\n".join(current_dict["gelisim_alanlari"]) or None,
                    gelecek_beklentileri="\n".join(current_dict["gelecek_beklentileri"]) or None
                )
                session.add(fb)
                
        # Buffer'ı sıfırla
        current_dict = {
            "harf_notu": "",
            "genel_degerlendirme": [],
            "guclu_alanlar": [],
            "gelisim_alanlari": [],
            "gelecek_beklentileri": []
        }
        current_category = None

    # 3. Paragraf Döngüsü ve Mantığı
    for line in paragraphs:
        lower_line = line.lower()
        clean_line = re.sub(r'^[\W_]+', '', line).strip() # Satır başındaki emoji/madde işaretlerini temizle
        clean_lower = clean_line.lower()

        # A. Başlık mı?
        matched_category = None
        for keyword, column_name in FEEDBACK_COLUMN_MAP.items():
            if keyword in lower_line:
                matched_category = column_name
                break
                
        if matched_category:
            current_category = matched_category
            continue

        # B. İsim mi?
        if len(clean_line) <= 45 and 0 < len(clean_line.split()) <= 4 and not clean_line.endswith('.') and not clean_line.endswith(';'):
            matches = difflib.get_close_matches(clean_lower, known_names, n=1, cutoff=0.8)
            if matches:
                flush_buffer()
                current_sicil = names_dict[matches[0]]
                current_year = None
                continue
            elif current_category is None and current_sicil is None and not clean_line.startswith("202"):
                print(f"❌ BULUNAMADI: {line}")

        # C. Yıl veya Harf Notu mu? (Başlık niteliğinde olan kısa satırlar)
        is_year_header = clean_line.startswith("202") and len(clean_line) <= 50
        is_harf_notu_line = "harf notu" in lower_line and ":" in line

        if is_year_header or is_harf_notu_line:
            if is_year_header:
                new_year = clean_line[0:4]
                if new_year != current_year:
                    flush_buffer()
                    current_year = new_year

            if is_harf_notu_line:
                current_dict["harf_notu"] = clean_line.split(":", 1)[1].strip() if ":" in clean_line else line.split(":", 1)[1].strip()
            
            # Başlık olduğu için içeriğe eklenmesin, devam (skip)
            continue
        
        # D. İçerik mi? (Yukarıdaki filtrelere takılmayan + null olmayan her şey içeriktir)
        if current_sicil and current_year and current_category and line:
            # Satır tek başına bir metinse madde işareti olarak ekle
            # Zaten Word'den gelen kendi madesi olabilir, - ekleyelim:
            bullet = "- " if not line.startswith("-") else ""
            current_dict[current_category].append(f"{bullet}{line}")

    # Döngü bitince elde kalan son tamponu da veritabanına yaz
    flush_buffer()

    # 4. Veritabanı Kaydı
    session.commit()

    count = session.query(EmployeeFeedback).count()
    print(f"EmployeeFeedback: {count} geri bildirim kaydedildi.")


# ──────────────────────────────────────────────────────────────────────────────
# Ana Seed Fonksiyonu (Değiştirilmedi — yalnızca yeni çağrılar eklendi)
# ──────────────────────────────────────────────────────────────────────────────

def run_seed():
    print("Veritabanı oluşturuluyor...")
    engine = create_engine(Config.DATABASE_URL)
    Base.metadata.drop_all(engine)
    Base.metadata.create_all(engine)
    Session = sessionmaker(bind=engine)
    session = Session()

    print("Excel verisi yükleniyor...")
    excel_path = os.path.join(Config.DATA_DIR, 'Verileri_SMART Hedefler ve Gerçekleşmeler.xlsx')
    if not os.path.exists(excel_path):
         print(f"HATA: {excel_path} bulunamadı.")
         return

    # Pandas okuma ve NaN temizliği
    df = pd.read_excel(excel_path)
    df.columns = [str(c).strip() for c in df.columns]
    
    # NaN floatları 0.0, stringleri 'Belirtilmemiş' yapalım
    for col in df.columns:
        if df[col].dtype == 'float64' or df[col].dtype == 'int64':
             df[col] = df[col].fillna(0.0)
        else:
             df[col] = df[col].fillna("Belirtilmemiş")

    # Benzersiz çalışanları bul
    employees_df = df[['Sicil', 'İsim', 'Unvan', 'Bölüm Ana Sorumluluk Alanı']].drop_duplicates()

    default_password = generate_password_hash('123456')

    # 1. Kök Düğüm (Üretim Operasyonları Müdürü) Create
    root_user = User(sicil_no='ROOT01', password_hash=default_password, role='Admin', manager_sicil=None)
    session.add(root_user)
    session.flush()
    
    root_emp = Employee(
        user_sicil='ROOT01',
        first_name='TUSAŞ',
        last_name='Üretim Operasyonları Müdürü',
        title='Üretim Operasyonları Müdürü',
        department='Üretim Yönetimi',
        email='root01@firma.com'
    )
    session.add(root_emp)
    session.commit()
    
    # 2. Her departman için bir Yönetici (Manager) Belirle
    dept_manager_map = {} # dept -> sicil_no mapped
    employee_id_map = {} # sicil -> employee_id to link goals

    # Tüm eşsiz departmanları alılım
    departments = employees_df['Bölüm Ana Sorumluluk Alanı'].unique()
    
    for dept in departments:
        dept_emps = employees_df[employees_df['Bölüm Ana Sorumluluk Alanı'] == dept]
        
        # Bu departmanda Müdür, Kıdemli veya en üst rütbeli birini Yönetici (Manager) seçelim
        manager_row = None
        for unvan in ['Müdür', 'Kıdemli Mühendis', 'Kıdemli Uzman', 'Uzman', 'Mühendis', 'Teknisyen']:
            match = dept_emps[dept_emps['Unvan'] == unvan]
            if not match.empty:
                manager_row = match.iloc[0]
                break
                
        if manager_row is not None:
            m_sicil = str(manager_row['Sicil']).strip()
            # Onu manager olarak tanımla
            manager_user = session.query(User).filter_by(sicil_no=m_sicil).first()
            if not manager_user:
                manager_user = User(
                    sicil_no=m_sicil, 
                    password_hash=default_password,
                    role='Manager', 
                    manager_sicil='ROOT01'
                )
                session.add(manager_user)
                session.flush()
                
                emp = Employee(
                    user_sicil=m_sicil,
                    first_name=str(manager_row['İsim']).split(" ")[0],
                    last_name=" ".join(str(manager_row['İsim']).split(" ")[1:]) if len(str(manager_row['İsim']).split(" ")) > 1 else "",
                    title=manager_row['Unvan'] + " / Bölüm Sorumlusu",  # Etiket ekleyelim
                    department=dept,
                    email=f"{m_sicil}@firma.com"
                )
                session.add(emp)
                session.flush()
            
            dept_manager_map[dept] = m_sicil
            employee_id_map[m_sicil] = emp.id if emp else session.query(Employee).filter_by(user_sicil=m_sicil).first().id

    # 3. Kalan diğer personelleri bölümlerinin Manager'ına bağlayarak (Employee) oluştur
    for idx, row in employees_df.iterrows():
        sicil = str(row['Sicil']).strip()
        dept = row['Bölüm Ana Sorumluluk Alanı']
        
        # Eğer bu sicil dept_manager_map içinde değer olarak varsa zaten Manager olarak üretilmiştir.
        if sicil in dept_manager_map.values():
            continue 
        
        m_sicil = dept_manager_map.get(dept, 'ROOT01') 
        
        user = User(
            sicil_no=sicil,
            password_hash=default_password,
            role='Employee',
            manager_sicil=m_sicil
        )
        session.add(user)
        session.flush()
        
        emp = Employee(
            user_sicil=sicil,
            first_name=str(row['İsim']).split(" ")[0],
            last_name=" ".join(str(row['İsim']).split(" ")[1:]) if len(str(row['İsim']).split(" ")) > 1 else "",
            title=row['Unvan'],
            department=dept,
            email=f"{sicil}@firma.com"
        )
        session.add(emp)
        session.flush()
        employee_id_map[sicil] = emp.id

    session.commit()

    # 4. PerformanceHistory
    for idx, row in df.iterrows():
        sicil = str(row.get('Sicil No', row.get('Sicil', ''))).strip()
        e_id = employee_id_map.get(sicil)
        if not e_id:
             continue
        
        ph = PerformanceHistory(
             employee_id=e_id,
             sicil_no=str(row.get('Sicil', '')).strip(),
             isim=str(row.get('İsim', '')).strip(),
             bolum=str(row.get('Bölüm Ana Sorumluluk Alanı', '')).strip(),
             unvan=str(row.get('Unvan', '')).strip(),
             yil=str(row.get('Yıl', '')).replace(".0","").strip(),
             hedef_turu=str(row.get('Hedef Türü', '')).strip(),
             yetkinlik=str(row.get('Yetkinlik', '')).strip(),
             stratejik_hedef=str(row.get('Stratejik Hedef Tanımı', '')).strip(),
             smart_hedef=str(row.get('SMART Hedef Tanımı', '')).strip(),
             hedef_degeri=str(row.get('Hedef Değeri', '')).strip(),
             birim=str(row.get('Birim', '')).strip(),
             hedef_yonu=str(row.get('Hedef Yönü', '')).strip(),
             gerceklesen_deger=str(row.get('Gerçekleşen Değer', '')).strip(),
             sonuc=str(row.get('Gerçekleşen Değere Göre Sonuç', '')).strip()
        )
        session.add(ph)

    session.commit()
    print("İlişkisel Veritabanı (SQLite) aktarımı tamamlandı.")

    # ── YENİ: Görev Tanımları ve Geri Bildirimleri ayrıştır ──────────
    parse_job_descriptions(session, Config.DATA_DIR)
    parse_feedbacks(session, Config.DATA_DIR)

    # 5. ChromaDB Organizasyon Şeması Entegresi
    print("Word dosyası ayrıştırılıyor ve ChromaDB'ye ekleniyor...")
    word_path = os.path.join(Config.DATA_DIR, 'ORGANİZASYON ŞEMASI VE GÖREV TANIMLARI.docx')
    if not os.path.exists(word_path):
        print("Word belgesi bulunamadı, ChromaDB adımı atlanıyor.")
        return

    doc = Document(word_path)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", ".", " ", ""]
    )
    chunks = text_splitter.split_text(full_text)

    chroma_client = chromadb.PersistentClient(path=Config.CHROMA_DB_PATH)
    embedding_func = embedding_functions.SentenceTransformerEmbeddingFunction(
        model_name=Config.EMBEDDING_MODEL
    )
    collection = chroma_client.get_or_create_collection(
        name="pms_data",
        embedding_function=embedding_func
    )

    ids = []
    documents = []
    metadatas = []

    for i, c in enumerate(chunks):
        # Bu chunka dahil olan sicilleri bul
        related_sicils = []
        for sicil, r in employees_df.iterrows():
             unvan = str(r['Unvan']).lower()
             bolum = str(r['Bölüm Ana Sorumluluk Alanı']).lower()
             
             # Eğer chunk içinde kişinin unvanı geçiyorsa sicili ekle
             if unvan in c.lower() or bolum in c.lower():
                  # string olan sicili kullan
                  related_sicils.append(str(r['Sicil']).strip())
        
        # Çok fazla sicil varsa ilk 10 tanesini alalım veya string olarak virgülle birleştirelim
        related_sicils = list(set(related_sicils))
        sicil_str = ",".join(related_sicils) if related_sicils else "Genel"

        ids.append(f"org_docs_{i}_{os.urandom(4).hex()}")
        documents.append(c)
        metadatas.append({
             "source": "ORGANİZASYON ŞEMASI VE GÖREV TANIMLARI.docx",
             "related_sicils": sicil_str
        })

    collection.add(
        ids=ids,
        documents=documents,
        metadatas=metadatas
    )
    print(f"ChromaDB'ye {len(documents)} parça (ilgili sicil metadatalarıyla) aktarıldı.")

if __name__ == "__main__":
    run_seed()
